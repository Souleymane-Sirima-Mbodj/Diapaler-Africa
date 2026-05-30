"""
DIAPALER AFRICA - Generateur DOCX academique (sans couleur)
============================================================
Usage : python generer_docx.py
Sortie : dossier docx/ avec 6 fichiers DOCX prêts à remettre.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONFIG GROUPE
# ─────────────────────────────────────────────
MEMBRES = [
    "Alioune Badara Barry",
    "Anta Diama Kama",
    "Souleymane Sirima Mbodj",
    "Serigne Abdoul Aziz Ndiaye",
    "Mohamed Moctar Niang",
    "Mareme Tine",
]

CLASSE      = "L3 Genie Informatique"    # <- modifier si besoin
PROFESSEUR  = "[Nom du Professeur]"      # <- modifier si besoin
DATE_REMISE = "24 mai 2026"

LIVRABLES = [
    (1, "Creation du projet Flutter, Interfaces, Navigation"),
    (2, "Consommation d'API Externes"),
    (3, "Authentification"),
    (4, "Gestion des Profils"),
    (5, "Fonctionnalites Avancees"),
    (6, "Rapport Final de Projet"),
]

NOIR  = RGBColor(0x00, 0x00, 0x00)
GRIS  = RGBColor(0x60, 0x60, 0x60)
GRIS_CLAIR_HEX = "F2F2F2"
BLANC_HEX      = "FFFFFF"


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_shading(cell, hex_fill: str):
    """Fond de cellule en hex sans '#'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Filet horizontal fin."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def page_setup(doc):
    """Marges A4 standard academique."""
    s = doc.sections[0]
    s.page_width    = Cm(21.0)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.0)


def add_header_footer(doc, numero: str, titre: str):
    """En-tête et pied de page discrets."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # En-tete (pages 2+)
    hdr = section.header
    hdr.paragraphs[0].clear()
    p = hdr.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"DIAPALER AFRICA  |  {numero} - {titre}")
    r1.font.size = Pt(9)
    r1.font.name = 'Times New Roman'
    r1.font.color.rgb = GRIS
    r1.italic = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Pied de page
    ftr = section.footer
    ftr.paragraphs[0].clear()
    pf = ftr.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(2)
    rf = pf.add_run("Ecole Superieure Polytechnique de Dakar  |  2025-2026  |  Page ")
    rf.font.size = Pt(8)
    rf.font.name = 'Times New Roman'
    rf.font.color.rgb = GRIS
    # Numero de page automatique
    run_el = pf.add_run()._r
    for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
        el = OxmlElement(tag)
        if text:
            el.text = text
        else:
            el.set(qn('w:fldCharType'), 'begin' if run_el.getparent() is None else 'end')
        run_el.append(el)
    # fix: recreer proprement
    pf2 = ftr.paragraphs[0]
    pf2.clear()
    pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_label = pf2.add_run("Ecole Superieure Polytechnique de Dakar  |  2025-2026  |  Page ")
    r_label.font.size = Pt(8)
    r_label.font.name = 'Times New Roman'
    r_label.font.color.rgb = GRIS
    # champ PAGE
    r_page = pf2.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r_page._r.extend([fldChar1, instr, fldChar2])


# ─────────────────────────────────────────────
# PAGE DE GARDE CLASSIQUE
# ─────────────────────────────────────────────

def add_cover_page(doc, num: int, titre: str):
    """Page de garde sobre et académique."""

    # Espacement haut
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Institution
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_u = p_univ.add_run("ECOLE SUPERIEURE POLYTECHNIQUE DE DAKAR")
    r_u.bold = True
    r_u.font.size = Pt(13)
    r_u.font.name = 'Times New Roman'
    r_u.font.color.rgb = NOIR

    p_dep = doc.add_paragraph()
    p_dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_d = p_dep.add_run("Departement Genie Informatique")
    r_d.font.size = Pt(11)
    r_d.font.name = 'Times New Roman'

    add_horizontal_rule(doc)

    # Espacement
    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Titre projet
    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_proj = p_proj.add_run("DIAPALER AFRICA")
    r_proj.bold = True
    r_proj.font.size = Pt(22)
    r_proj.font.name = 'Times New Roman'
    r_proj.font.color.rgb = NOIR

    p_sous = doc.add_paragraph()
    p_sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sous.add_run("Plateforme mobile de mentorat entrepreneurial au Senegal")
    r_s.italic = True
    r_s.font.size = Pt(12)
    r_s.font.name = 'Times New Roman'

    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    add_horizontal_rule(doc)

    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Numero et titre du livrable
    p_liv = doc.add_paragraph()
    p_liv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ln = p_liv.add_run(f"LIVRABLE {num}")
    r_ln.bold = True
    r_ln.font.size = Pt(18)
    r_ln.font.name = 'Times New Roman'

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tt = p_tit.add_run(titre)
    r_tt.bold = True
    r_tt.font.size = Pt(14)
    r_tt.font.name = 'Times New Roman'

    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Tableau d'identification
    table = doc.add_table(rows=len(MEMBRES) + 4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def fill(idx, label, valeur, gras_label=True):
        row = table.rows[idx]
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)
        # Label
        pl = row.cells[0].paragraphs[0]
        pl.paragraph_format.space_before = Pt(3)
        pl.paragraph_format.space_after  = Pt(3)
        rl = pl.add_run(label)
        rl.font.name = 'Times New Roman'
        rl.font.size = Pt(10)
        rl.bold = gras_label
        # Valeur
        pv = row.cells[1].paragraphs[0]
        pv.paragraph_format.space_before = Pt(3)
        pv.paragraph_format.space_after  = Pt(3)
        rv = pv.add_run(valeur)
        rv.font.name = 'Times New Roman'
        rv.font.size = Pt(10)

    for i, m in enumerate(MEMBRES):
        fill(i, f"Membre {i+1}", m)

    fill(len(MEMBRES),     "Classe / Filiere",   CLASSE)
    fill(len(MEMBRES) + 1, "Enseignant",          PROFESSEUR)
    fill(len(MEMBRES) + 2, "Module",              "Developpement d'Applications Mobiles")
    fill(len(MEMBRES) + 3, "Date de remise",      DATE_REMISE)

    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Annee
    p_an = doc.add_paragraph()
    p_an.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_an = p_an.add_run("Annee academique 2025 - 2026")
    r_an.font.size = Pt(10)
    r_an.font.name = 'Times New Roman'
    r_an.italic = True

    # Saut de page
    doc.add_page_break()


# ─────────────────────────────────────────────
# INLINE MARKDOWN (gras, italique, code)
# ─────────────────────────────────────────────

def apply_inline(para, text: str):
    """Ajoute du texte avec gras/italique/code dans un paragraphe."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)', re.DOTALL)
    for tok in pattern.findall(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = para.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            r = para.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = para.add_run(tok[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            r = para.add_run(tok)
        r.font.name = r.font.name or 'Times New Roman'


def set_font(para, size=Pt(11), name='Times New Roman', bold=False, italic=False, color=None):
    for r in para.runs:
        r.font.size  = size
        r.font.name  = name
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = color


def add_screenshot_placeholder(doc, label: str):
    """Cadre gris adaptatif comme emplacement pour une capture d'écran.

    La cellule a une hauteur MINIMALE de 3 cm (atLeast) et s'agrandit
    automatiquement pour contenir l'image insérée — elle ne déborde jamais.
    """

    # Ligne de label (italique gris)
    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_before = Pt(6)
    p_lbl.paragraph_format.space_after  = Pt(2)
    r = p_lbl.add_run(f'📸  {label}')
    r.font.name      = 'Times New Roman'
    r.font.size      = Pt(9)
    r.font.color.rgb = GRIS
    r.italic         = True

    # Cadre placeholder : table 1×1, fond gris clair
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, 'F0F0F0')

    # Hauteur MINIMALE 3 cm (atLeast) — la cellule GRANDIT avec l'image
    # Contrairement à 'exact', 'atLeast' ne bloque pas le contenu qui dépasse.
    tr   = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '1701')    # 3 cm (1 cm ≈ 567 twips)
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

    # Marge interne de la cellule pour que l'image ne touche pas les bords
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    '113')   # ~2 mm
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

    # Texte d'indication centré (remplacé par l'image lors de l'insertion)
    pc = cell.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = Pt(10)
    pc.paragraph_format.space_after  = Pt(10)
    rc = pc.add_run('[ Insérer la capture d\'écran ici ]')
    rc.font.name      = 'Times New Roman'
    rc.font.size      = Pt(10)
    rc.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
    rc.italic         = True

    # Espace après le cadre
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────
# PARSEUR MARKDOWN -> DOCX
# ─────────────────────────────────────────────

def parse_md(doc, lines: list):
    i = 0
    in_code = False
    code_buf = []
    code_lang = ''
    table_buf = []

    while i < len(lines):
        line = lines[i]

        # ── Bloc de code ──────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code   = True
                code_lang = line.strip()[3:].strip()
                code_buf  = []
                i += 1
                continue
            else:
                in_code = False
                # Étiquette langue
                if code_lang:
                    pl = doc.add_paragraph()
                    pl.paragraph_format.space_after = Pt(0)
                    rl = pl.add_run(code_lang.upper())
                    rl.bold = True
                    rl.font.size = Pt(8)
                    rl.font.name = 'Times New Roman'

                # Tableau 1x1 fond gris très clair
                tbl = doc.add_table(rows=1, cols=1)
                tbl.style = 'Table Grid'
                cell = tbl.rows[0].cells[0]
                set_cell_shading(cell, GRIS_CLAIR_HEX)
                # Vider le 1er para vide
                if cell.paragraphs and not cell.paragraphs[0].text:
                    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                for cl in code_buf:
                    pc = cell.add_paragraph(cl)
                    pc.paragraph_format.space_before = Pt(0)
                    pc.paragraph_format.space_after  = Pt(0)
                    for r in pc.runs:
                        r.font.name = 'Courier New'
                        r.font.size = Pt(8.5)

                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)
                code_buf  = []
                code_lang = ''
                i += 1
                continue

        if in_code:
            code_buf.append(line.rstrip())
            i += 1
            continue

        # ── Tables Markdown ───────────────────
        if line.strip().startswith('|'):
            table_buf.append([c.strip() for c in line.strip().strip('|').split('|')])
            i += 1
            continue
        elif table_buf:
            rows_clean = [r for r in table_buf
                          if not all(re.match(r'^-+$', c.replace(':', '')) for c in r if c)]
            if rows_clean:
                nb_cols = max(len(r) for r in rows_clean)
                tbl = doc.add_table(rows=len(rows_clean), cols=nb_cols)
                tbl.style = 'Table Grid'
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                for ri, rdata in enumerate(rows_clean):
                    is_header = (ri == 0)
                    for ci in range(nb_cols):
                        val = rdata[ci] if ci < len(rdata) else ''
                        cell = tbl.rows[ri].cells[ci]
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        if is_header:
                            set_cell_shading(cell, 'D9D9D9')  # gris header
                        pc = cell.paragraphs[0]
                        pc.paragraph_format.space_before = Pt(2)
                        pc.paragraph_format.space_after  = Pt(2)
                        apply_inline(pc, val)
                        set_font(pc, size=Pt(9.5), bold=is_header)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)
            table_buf = []

        stripped = line.strip()

        # ── Ignorer éléments page de garde ────
        if not stripped or stripped == '&nbsp;' or stripped.startswith('<!--'):
            i += 1
            continue

        if stripped == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Titres ────────────────────────────
        if stripped.startswith('#### '):
            p = doc.add_paragraph(style='Heading 4')
            apply_inline(p, stripped[5:])
            set_font(p, size=Pt(11), bold=True)
            i += 1; continue

        if stripped.startswith('### '):
            p = doc.add_paragraph(style='Heading 3')
            apply_inline(p, stripped[4:])
            set_font(p, size=Pt(12), bold=True)
            i += 1; continue

        if stripped.startswith('## '):
            p = doc.add_paragraph(style='Heading 2')
            apply_inline(p, stripped[3:])
            set_font(p, size=Pt(13), bold=True)
            i += 1; continue

        if stripped.startswith('# '):
            txt = re.sub(r'!\[.*?\]', '', stripped[2:]).strip()
            p = doc.add_paragraph(style='Heading 1')
            apply_inline(p, txt)
            set_font(p, size=Pt(15), bold=True)
            i += 1; continue

        # ── Citations (>) ─────────────────────
        if stripped.startswith('> '):
            content = stripped[2:]

            # Placeholder capture d'écran : ligne contenant 📸
            if '📸' in content:
                # Extraire le label : supprimer les marqueurs ** et le préfixe "📸 CAPTURE D'ÉCRAN — "
                clean = re.sub(r'\*+', '', content)
                clean = re.sub(r'`[^`]*`', lambda m: m.group(0)[1:-1], clean)  # garder texte des backticks
                m_lbl = re.search(r'📸[^—–]*[—–]+\s*(.*)', clean)
                label = m_lbl.group(1).strip() if m_lbl else re.sub(r'📸\s*', '', clean).strip()
                label = label.strip('* ').strip()
                add_screenshot_placeholder(doc, label or 'Capture d\'écran')
                i += 1
                # Sauter la ligne "*(Insérer ici la capture d'écran)*" qui suit
                if i < len(lines) and 'Insérer' in lines[i]:
                    i += 1
                continue

            # Citation normale
            p = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            apply_inline(p, content)
            p.paragraph_format.left_indent = Cm(1)
            set_font(p, size=Pt(10), italic=True, color=GRIS)
            i += 1; continue

        # ── Listes ────────────────────────────
        if re.match(r'^[-*]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet')
            apply_inline(p, re.sub(r'^[-*]\s+', '', stripped))
            set_font(p, size=Pt(10.5))
            i += 1; continue

        if re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph(style='List Number')
            apply_inline(p, re.sub(r'^\d+\.\s+', '', stripped))
            set_font(p, size=Pt(10.5))
            i += 1; continue

        # ── Paragraphe normal ─────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        apply_inline(p, stripped)
        set_font(p, size=Pt(11))

        i += 1


# ─────────────────────────────────────────────
# GENERATION
# ─────────────────────────────────────────────

def generate(md_path: Path, out_path: Path, num: int, titre: str):
    print(f"  Livrable {num} ... ", end="", flush=True)

    doc = Document()
    page_setup(doc)

    # Style Normal de base
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    add_header_footer(doc, f"Livrable {num}", titre)
    add_cover_page(doc, num, titre)

    # Lire le MD et trouver le debut du contenu (apres la page de garde)
    text = md_path.read_text(encoding='utf-8')
    m = re.search(r'\n# LIVRABLE \d+', text)
    content = text[m.start():] if m else '\n'.join(text.split('\n')[50:])

    parse_md(doc, content.split('\n'))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print("OK")


def main():
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    base = Path(__file__).parent
    out  = base / "docx"
    out.mkdir(exist_ok=True)

    print("\n=== DIAPALER AFRICA - Generation DOCX academique ===\n")
    for num, titre in LIVRABLES:
        md = base / f"livrable{num}.md"
        if not md.exists():
            print(f"  [ERREUR] {md.name} introuvable")
            continue
        generate(md, out / f"DIAPALER_AFRICA_Livrable{num}.docx", num, titre)

    print(f"\nFichiers generes dans : {out.resolve()}")
    print("\nA faire :")
    print("  1. Remplacer [Nom du Professeur] et [Ta Classe] si besoin")
    print("  2. Inserer les captures d'ecran dans Word")


if __name__ == "__main__":
    main()
