"""
Générateur de livrables DIAPALER AFRICA : .md → .docx
Utilise python-docx pour produire des fichiers Word professionnels.
"""

import re
import sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

# ── Chemins ──────────────────────────────────────────────────────────────────
BASE = Path(r"C:\Users\HP\OneDrive\Documents\Claude\Projects\ndiaye\livrables")
DOCX_DIR = BASE / "docx"
DOCX_DIR.mkdir(exist_ok=True)

# ── Couleurs ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
BLUE   = RGBColor(0x1A, 0x56, 0xDB)
GREEN  = RGBColor(0x05, 0x71, 0x3A)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
RED    = RGBColor(0xC8, 0x1E, 0x1E)
GREY   = RGBColor(0x37, 0x41, 0x51)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF3, 0xF4, 0xF6)   # fond code
DGREY  = RGBColor(0x6B, 0x72, 0x80)   # texte muted

# ── Helpers XML ───────────────────────────────────────────────────────────────

def set_cell_shading(cell, fill_hex):
    """Colore le fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_para_shading(para, fill_hex):
    """Colore le fond d'un paragraphe (pour les blocs de code)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_para_border(para, color_hex='CCCCCC', sz='4'):
    """Ajoute une bordure gauche à un paragraphe (pour les blockquotes)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return para

# ── Styles de base ─────────────────────────────────────────────────────────────

def apply_heading_style(para, level):
    """Applique le style heading avec la bonne couleur."""
    colors = {1: NAVY, 2: NAVY, 3: RGBColor(0x1E, 0x3A, 0x8A)}
    for run in para.runs:
        run.font.color.rgb = colors.get(level, NAVY)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt({1: 18, 2: 15, 3: 13}.get(level, 12))

def configure_doc_styles(doc):
    """Configure les styles du document."""
    # Style Normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Style Heading 1
    for level, name in [(1,'Heading 1'), (2,'Heading 2'), (3,'Heading 3')]:
        try:
            h = doc.styles[name]
            h.font.name = 'Calibri'
            h.font.bold = True
            h.font.color.rgb = NAVY
            h.font.size = Pt({1:18, 2:15, 3:13}[level])
            h.paragraph_format.space_before = Pt({1:18, 2:14, 3:10}[level])
            h.paragraph_format.space_after  = Pt({1:8, 2:6, 3:4}[level])
        except Exception:
            pass

# ── Page de garde ─────────────────────────────────────────────────────────────

def add_title_page(doc, livrable_num, livrable_title, livrable_subtitle):
    """Génère la page de garde DIAPALER AFRICA."""
    # En-tête institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ÉCOLE SUPÉRIEURE POLYTECHNIQUE DE DAKAR')
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Département Génie Informatique')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.font.color.rgb = GREY

    # Séparateur
    for _ in range(3):
        doc.add_paragraph()

    # Titre projet
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('DIAPALER AFRICA')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(28)
    run3.font.bold = True
    run3.font.color.rgb = NAVY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Plateforme mobile de mentorat entrepreneurial au Sénégal')
    run4.font.name = 'Calibri'
    run4.font.size = Pt(13)
    run4.font.color.rgb = GREY

    # Séparateur
    for _ in range(2):
        doc.add_paragraph()

    # Numéro du livrable
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(f'LIVRABLE {livrable_num}')
    run5.font.name = 'Calibri'
    run5.font.size = Pt(22)
    run5.font.bold = True
    run5.font.color.rgb = AMBER

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run(livrable_title)
    run6.font.name = 'Calibri'
    run6.font.size = Pt(16)
    run6.font.bold = True
    run6.font.color.rgb = NAVY

    if livrable_subtitle:
        p7 = doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run7 = p7.add_run(livrable_subtitle)
        run7.font.name = 'Calibri'
        run7.font.size = Pt(12)
        run7.font.color.rgb = GREY

    for _ in range(3):
        doc.add_paragraph()

    # Tableau membres
    membres = [
        ('Membre 1', 'Alioune Badara Barry'),
        ('Membre 2', 'Anta Diama Kama'),
        ('Membre 3', 'Souleymane Sirima Mbodj'),
        ('Membre 4', 'Serigne Abdoul Aziz Ndiaye'),
        ('Membre 5', 'Mohamed Moctar Niang'),
        ('Membre 6', 'Mareme Tine'),
        ('Classe / Filière', '[Ta Classe]'),
        ('Enseignant', '[Nom du Professeur]'),
        ('Module', 'Développement d\'Applications Mobiles'),
        ('Institution', 'École Supérieure Polytechnique (ESP) — Dakar'),
        ('Année académique', '2025 – 2026'),
        ('Date de remise', '[Date]'),
    ]
    table = doc.add_table(rows=len(membres), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(membres):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], 'EFF6FF')

    doc.add_paragraph()

    # Année académique
    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_y = p_year.add_run('Année académique : 2025 – 2026')
    run_y.font.name = 'Calibri'
    run_y.font.size = Pt(11)
    run_y.font.color.rgb = GREY

    # Page break après la page de garde
    add_page_break(doc)

# ── Rendu inline Markdown ──────────────────────────────────────────────────────

def add_inline_text(run_parent, text):
    """
    Parse le texte avec **bold**, *italic*, `code` et ajoute les runs correspondants.
    run_parent est un objet paragraph.
    """
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # Texte avant
        if m.start() > pos:
            run = run_parent.add_run(text[pos:m.start()])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        full = m.group(0)
        if full.startswith('**'):
            run = run_parent.add_run(m.group(2))
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif full.startswith('*'):
            run = run_parent.add_run(m.group(3))
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif full.startswith('`'):
            run = run_parent.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        pos = m.end()

    # Reste
    if pos < len(text):
        run = run_parent.add_run(text[pos:])
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

# ── Blocs de code ──────────────────────────────────────────────────────────────

def add_code_block(doc, code_lines):
    """Ajoute un bloc de code formaté (fond gris, Courier New)."""
    # Paragraphe d'en-tête du bloc (juste un séparateur visuel)
    for line in code_lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.5)
        set_para_shading(para, 'F3F4F6')

        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)

    # Espace après le bloc
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Tableaux Markdown ──────────────────────────────────────────────────────────

def add_md_table(doc, rows_data, header_row=True):
    """Crée un tableau Word depuis des données de tableau markdown."""
    if not rows_data:
        return

    ncols = len(rows_data[0])
    nrows = len(rows_data)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            para = cell.paragraphs[0]
            para.clear()
            add_inline_text(para, cell_text.strip())
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
            if i == 0 and header_row:
                for run in para.runs:
                    run.bold = True
                set_cell_shading(cell, '1E3A8A')
                for run in para.runs:
                    run.font.color.rgb = WHITE
            elif i % 2 == 1:
                set_cell_shading(cell, 'EFF6FF')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def parse_md_table(lines, start_idx):
    """Parse un tableau markdown et retourne (rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Ligne de séparation |---|---|
        if re.match(r'^\|[-:| ]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i

# ── Screenshot placeholders ────────────────────────────────────────────────────

def add_screenshot_placeholder(doc, text):
    """Ajoute un placeholder de capture d'écran : label amber + boîte bordée 4 côtés."""
    # ── Label amber ──────────────────────────────────────────────────────────────
    label = doc.add_paragraph()
    label.paragraph_format.left_indent  = Cm(0.0)
    label.paragraph_format.space_before = Pt(10)
    label.paragraph_format.space_after  = Pt(0)
    set_para_shading(label, 'FEF3C7')
    set_para_border(label, 'F59E0B', '8')
    run = label.add_run('📸  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x09)

    # ── Boîte rectangulaire (tableau 1×1) ────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'

    cell = tbl.cell(0, 0)
    set_cell_shading(cell, 'F9FAFB')

    # Hauteur minimale : 7 cm = 3969 twips — s'agrandit automatiquement si image insérée
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '3969')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Texte centré dans la boîte
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Cm(2.8)
    inner = para.add_run('[ Insérer la capture d\'écran ici ]')
    inner.font.name = 'Calibri'
    inner.font.size = Pt(9)
    inner.font.color.rgb = RGBColor(0xAB, 0xB5, 0xBF)
    inner.italic = True

    # Espace après
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ── Parseur Markdown principal ─────────────────────────────────────────────────

def parse_and_render_md(doc, md_content):
    """Parse le contenu markdown et le rend dans le document Word."""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''

    # Sauter la page de garde dans le markdown (jusqu'au 2ème ---)
    dash_count = 0
    while i < len(lines):
        if lines[i].strip() == '---':
            dash_count += 1
            if dash_count >= 2:
                i += 1
                break
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Bloc de code (début/fin) ──────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # ── Sauter les lignes vides et séparateurs ────────────────────────────
        if not stripped or stripped == '---' or stripped == '&nbsp;':
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith('#### '):
            para = doc.add_heading(stripped[5:], level=3)
            i += 1
            continue

        if stripped.startswith('### '):
            para = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith('## '):
            para = doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith('# '):
            para = doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # ── Tableau Markdown ──────────────────────────────────────────────────
        if stripped.startswith('|'):
            rows, next_i = parse_md_table(lines, i)
            if rows:
                add_md_table(doc, rows)
            i = next_i
            continue

        # ── Screenshot placeholder ────────────────────────────────────────────
        if stripped.startswith('> **📸') or stripped.startswith('> 📸'):
            # Nettoyer le texte
            text = stripped.lstrip('> ').strip('*').strip()
            # Sauter la ligne suivante si c'est *(Insérer...)*
            add_screenshot_placeholder(doc, text)
            i += 1
            if i < len(lines) and '(Insérer ici' in lines[i]:
                i += 1
            continue

        # ── Sauter les lignes de capture "*Insérer*"
        if '*(Insérer ici la capture' in stripped:
            i += 1
            continue

        # ── Blockquote (autre) ────────────────────────────────────────────────
        if stripped.startswith('> '):
            text = stripped[2:].strip().strip('*')
            if '📸' in text or 'CAPTURE' in text:
                add_screenshot_placeholder(doc, text)
            else:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(0.8)
                set_para_border(para, '4B83F0', '6')
                add_inline_text(para, text)
            i += 1
            continue

        # ── Listes à puces ────────────────────────────────────────────────────
        if stripped.startswith('- ') or stripped.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            add_inline_text(para, stripped[2:])
            for run in para.runs:
                run.font.size = Pt(11)
            i += 1
            continue

        # ── Listes numérotées ─────────────────────────────────────────────────
        if re.match(r'^\d+\.\s', stripped):
            para = doc.add_paragraph(style='List Number')
            text_part = re.sub(r'^\d+\.\s', '', stripped)
            add_inline_text(para, text_part)
            i += 1
            continue

        # ── Ligne normale ─────────────────────────────────────────────────────
        # Sauter les meta du header md (Projet :, Module :, etc. déjà dans titre)
        if stripped.startswith('**Projet :**') or stripped.startswith('**Module :**') or \
           stripped.startswith('**Institution :**') or stripped.startswith('**Année académique :**'):
            i += 1
            continue

        # Sauter les ancres [lien](#ancre)
        if re.match(r'^\[.+\]\(#.+\)$', stripped):
            i += 1
            continue

        # Texte normal
        para = doc.add_paragraph()
        add_inline_text(para, stripped)
        i += 1

# ── Générateur principal ────────────────────────────────────────────────────────

LIVRABLES_META = {
    1: ('Architecture Flutter et Navigation', ''),
    2: ('Consommation d\'API Externes', ''),
    3: ('Authentification', ''),
    4: ('Gestion des Profils', ''),
    5: ('Fonctionnalités Avancées', ''),
    6: ('Rapport Final de Projet', 'Bilan complet — DIAPALER AFRICA'),
}

def generate_livrable(num):
    md_path  = BASE / f'livrable{num}.md'
    out_path = DOCX_DIR / f'DIAPALER_AFRICA_Livrable{num}.docx'

    if not md_path.exists():
        print(f'  [SKIP] {md_path} introuvable')
        return

    print(f'  Génération Livrable {num}...', end=' ')

    doc = Document()
    configure_doc_styles(doc)

    # Marges de la page
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    title, subtitle = LIVRABLES_META[num]
    add_title_page(doc, num, title, subtitle)

    md_content = md_path.read_text(encoding='utf-8')
    parse_and_render_md(doc, md_content)

    doc.save(str(out_path))
    print(f'OK → {out_path.name}')

# ── Point d'entrée ─────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Génération des livrables DIAPALER AFRICA (.md → .docx)\n')
    for num in [1, 2, 3, 4, 5, 6]:
        generate_livrable(num)
    print('\nTerminé !')
